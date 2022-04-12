from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import os
import xlrd

data_path = 'vas-data'
group_label_path = 'group_label_auto.xls'
visualization_dir = 'group_vis'

if not os.path.isdir(visualization_dir):
    os.makedirs(visualization_dir)


def get_file_text(file_path):
    text_file = []
    with open(file_path, 'r') as f:
        lines = f.readlines()
        for idx, line in enumerate(lines):
            text = line.strip().replace('	', ' ')
            if line.startswith('*') or line.startswith('%'):
                text = text.split('')[0]
                text_file.append(text)
    return text_file


def get_group_label():
    workbook = xlrd.open_workbook(group_label_path)
    worksheet = workbook.sheet_by_index(0)
    first_col = []
    for row in range(0, worksheet.nrows):
        v = worksheet.cell_value(row, 0)
        first_col.append(v)

    def get_xlsx_col_contents_single(col_num):
        elm_to_get = {}
        for row in range(1, worksheet.nrows):
            elm_to_get[first_col[row]] = set()
            v = str(worksheet.cell_value(row, col_num))
            if v != '':
                for single_range in v.split(','):
                    single_range = single_range.strip()
                    if single_range.isnumeric():
                        elm_to_get[first_col[row]].add(int(single_range) - 1)
                    else:
                        for i in range(int(single_range.split('-')[0]) - 1, int(single_range.split('-')[1])):
                            elm_to_get[first_col[row]].add(i)
            elm_to_get[first_col[row]] = list(elm_to_get[first_col[row]])
        print(elm_to_get)
        return elm_to_get

    return [get_xlsx_col_contents_single(i) for i in range(1, 8)]


def apply_visualization(group_labels, text_dict):
    for key, text_list in text_dict.items():
        print(key)
        line_group_label_list = [[] for _ in range(len(text_list))]
        for group_idx, group_label in enumerate(group_labels):
            key_group_label = group_label[key]
            line_idx = 0
            par_line_idx = 0
            while line_idx < len(text_list):
                line = text_list[line_idx]
                if line.startswith('*'):
                    if par_line_idx in key_group_label:
                        line_group_label_list[line_idx].append(group_idx)
                    line_idx += 1
                    while line_idx < len(text_list) and text_list[line_idx].startswith('%'):
                        if par_line_idx in key_group_label:
                            line_group_label_list[line_idx].append(group_idx)
                        line_idx += 1
                    par_line_idx += 1
                    line_idx -= 1
                line_idx += 1
        generate_docx(key, text_list, line_group_label_list)


def generate_docx(key, text_list, line_group_label_list):
    document = Document()

    document.add_heading(key, 0)

    label_name_list = ['Question', 'Music', 'Reminder/Alarm/Timer/List', 'Phone Call', 'Smart Home', 'Conversation',
                       'Call']

    for label_idx, label_name in enumerate(label_name_list):
        paragraph = document.add_paragraph()
        paragraph.add_run('#').font.highlight_color = label_idx + 1
        paragraph.add_run(label_name)

    paragraph = document.add_paragraph()
    paragraph.add_run('-------------')

    par_idx = 0

    for idx, text in enumerate(text_list):
        paragraph = document.add_paragraph()
        line_group_label = line_group_label_list[idx]

        for label in line_group_label:
            paragraph.add_run('#').font.highlight_color = label + 1
        if text.startswith('*'):
            par_idx += 1
            paragraph.add_run(str(par_idx) + ': ')
        paragraph.add_run(text)

    document.add_page_break()

    document.save(os.path.join(visualization_dir, key + '.docx'))


text_dict = dict()
for text_file_name in sorted(os.listdir(data_path)):
    if text_file_name.endswith('.cha'):
        text_file_path = os.path.join(data_path, text_file_name)
        text = get_file_text(text_file_path)
        text_dict[text_file_name.split('.')[0]] = text
        print(text_file_path)
        print(text)

group_labels = get_group_label()
print(group_labels)

apply_visualization(group_labels, text_dict)
